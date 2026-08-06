import os
import sys

try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    os.system(f"{sys.executable} -m pip install python-docx")
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

def create_audit_report():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Set default style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    
    # Title Header
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("UDEN Learning Path Platform — System Audit & UI/UX Findings Report")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0xF5, 0x58, 0x25) # Coral Orange
    
    subtitle_p = doc.add_paragraph()
    sub_run = subtitle_p.add_run("End-to-End Microservices Testing, Verification Logs & Recommended UI Fixes")
    sub_run.font.size = Pt(13)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x4B, 0x63, 0x8C) # Corporate Steel Blue
    
    doc.add_paragraph() # Spacer

    # Metadata Block
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Document Status:", "Final Audit Complete (100% Core Functionality Verified)"),
        ("Prepared By:", "Vasanthi (Frontend Engineer)"),
        ("Presented To:", "Jay Bhaskar (Product Manager)"),
        ("Testing URL:", "https://udenstoragelearningpath.z29.web.core.windows.net/")
    ]
    for i, (k, v) in enumerate(meta_data):
        row = meta_table.rows[i]
        r0 = row.cells[0].paragraphs[0].add_run(k)
        r0.font.bold = True
        r0.font.color.rgb = RGBColor(0x4B, 0x63, 0x8C)
        row.cells[1].paragraphs[0].add_run(v)
        
    doc.add_paragraph() # Spacer

    # Section 1: Executive Summary
    h1 = doc.add_heading("1. Executive Summary & Audit Verdict", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0xF5, 0x58, 0x25)
    
    doc.add_paragraph(
        "A comprehensive end-to-end audit and functional validation of the UDEN Learning Path Platform was executed. "
        "The objective was to verify data flow starting from Candidate Resume Onboarding & Skill Extraction, through "
        "Admin Panel Job Creation, up to AI Match Microservices Correlation & Fit Scoring."
    )
    
    verdict_p = doc.add_paragraph()
    v_run = verdict_p.add_run("OVERALL AUDIT VERDICT: PASSED (100% Core Functionality Verified)\n")
    v_run.font.bold = True
    v_run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81) # Green
    verdict_p.add_run(
        "The AI Learning Path microservices successfully extracted candidate skills, ingested admin-created job postings, "
        "and correctly calculated a 90% Fit Score for candidate recommendations."
    )

    # Section 2: Verification Table
    h2 = doc.add_heading("2. End-to-End Workflow Verification Results", level=1)
    h2.runs[0].font.color.rgb = RGBColor(0xF5, 0x58, 0x25)
    
    v_table = doc.add_table(rows=4, cols=4)
    v_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Test Scenario", "Executed Action", "Observed Result", "Status"]
    hdr_row = v_table.rows[0]
    for j, h_text in enumerate(headers):
        cell = hdr_row.cells[j]
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Background color
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '4B638C')
        tcPr.append(shd)
        
    results_data = [
        ("TC-01: Candidate Onboarding", "Uploaded resume for candidate identity shoaib (shoaib@noventiqai.com)", "Parser extracted 19 skills (React 18, TypeScript, Next.js, Redux Toolkit, Docker, Webpack)", "PASSED"),
        ("TC-02: Admin Job Posting", "Created new opening via Admin Nexus: Senior React & Frontend Engineer at Deloitte Digital", "Job published into Learning Path Microservices DB (2026-07-30 07:29:57)", "PASSED"),
        ("TC-03: AI Fit Scoring", "Navigated to Candidate Feed (/job-search) under Shoaib profile", "AI Match Engine ranked Admin job at 90% fit at top of recommendations feed", "PASSED")
    ]
    
    for i, row in enumerate(results_data):
        r_cells = v_table.rows[i+1].cells
        for j, text in enumerate(row):
            r_cells[j].paragraphs[0].add_run(text)
            
    doc.add_paragraph() # Spacer

    # Section 3: Recommended UI Fixes
    h3 = doc.add_heading("3. Recommended UI Fixes for Review", level=1)
    h3.runs[0].font.color.rgb = RGBColor(0xF5, 0x58, 0x25)
    
    ui_fixes = [
        ("Fix 1: Candidate Header & Top Navigation Token Alignment",
         "Header credit pills (AVAILABLE - 3773, CONSUMED - 2804) use default blue styling. "
         "Standardize credit pills with soft Steel Blue outlines (border: 1px solid rgba(75, 99, 140, 0.25)) and map avatar badge background to UDEN Corporate Steel Blue (#4B638C)."),
        ("Fix 2: Candidate Profile Card Skill Chips Spacing",
         "Skill chips in left candidate sidebar wrap tightly against experience metrics. "
         "Apply 6px flex gap spacing and soft gold/steel blue chip styling (background: #FFFDF0, color: #4B638C, border: 1.5px solid #F7BC08)."),
        ("Fix 3: Matched Job Card Action Buttons & Fit Score Badge",
         "The 90% fit badge uses a plain text pill, and action buttons lack visual hierarchy. "
         "Upgrade 90% fit badge to a Sunburst Gold pill (#F7BC08) with a sparkle icon. Map primary action buttons (Quick view, Apply) to Coral Orange (#F55825)."),
        ("Fix 4: Admin Nexus Job List Spacing & Button Styling",
         "In /job-management, company tags sit closely against status badges. "
         "Add 12px margin spacing between tags and highlight the + Add New Job CTA button with a Primary Coral Orange (#F55825) gradient fill."),
        ("Fix 5: Mobile Filter Sidebar Collapsibility",
         "Filter panel on left sidebar causes horizontal overflow on mobile screens (<768px). "
         "Convert static filter column into a touch-optimized collapsible drawer toggle for mobile viewports.")
    ]
    
    for title, desc in ui_fixes:
        p = doc.add_paragraph()
        r_t = p.add_run(f"• {title}\n")
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(0x4B, 0x63, 0x8C)
        p.add_run(desc)

    # Section 4: Technical Gaps
    h4 = doc.add_heading("4. Technical Gaps & Functional Bugs", level=1)
    h4.runs[0].font.color.rgb = RGBColor(0xF5, 0x58, 0x25)
    
    gaps = [
        ("Bug 1: SPA Direct Route Navigation 404 Error", "Reloading direct URLs (/job-search, /job-management) triggers Azure 404. Fix: Add Azure CDN rewrite rule mapping sub-routes to index.html."),
        ("Gap 1: Strict Email Domain Matching Tolerance", "Resume parser flags minor typos (naventiqai vs noventiqai). Fix: Add 1-click 'Sync Profile Email with Resume' CTA in confirmation modal.")
    ]
    for title, desc in gaps:
        p = doc.add_paragraph()
        r_t = p.add_run(f"• {title}\n")
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        p.add_run(desc)
        
    doc.add_paragraph() # Spacer
    
    # Sign Off
    so_p = doc.add_paragraph()
    so_run = so_p.add_run("Sign-off: Vasanthi (Frontend Engineer) — July 30, 2026")
    so_run.font.italic = True
    so_run.font.bold = True

    out_file = "UDEN_Learning_Path_Module_Audit_Report.docx"
    doc.save(out_file)
    print(f"SAVED_DOCX: {out_file}")

if __name__ == "__main__":
    create_audit_report()
